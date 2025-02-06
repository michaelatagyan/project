import io
import sys

from PyQt6 import uic
from PyQt6.QtWidgets import (QTextEdit, QApplication, QMainWindow, QFileDialog, QSizePolicy,
                             QVBoxLayout, QDialog, QPushButton, QHBoxLayout, QLabel)
from PyQt6.QtGui import QRegularExpressionValidator
from PyQt6.QtCore import QRegularExpression, Qt
from assets.template import template
from assets.possible_words import stress_dict
from docx import Document
from assets.task import tsk
import subprocess, os
from assets.task10_editor import Task10Editor
from assets.command import Command

class Main(QMainWindow):
    def __init__(self):
        super().__init__()
        f = io.StringIO(template)
        uic.loadUi(f, self)
        self.initUI()
        self.current_task = None

    def initUI(self):
        # Создаем основной QVBoxLayout для centralWidget
        main_layout = QVBoxLayout(self.centralwidget)
        main_layout.addWidget(self.horizontalLayoutWidget)
        main_layout.addWidget(self.horizontalLayoutWidget_2)
        main_layout.addWidget(self.horizontalLayoutWidget_3)
        main_layout.addWidget(self.horizontalLayoutWidget_4)
        main_layout.addWidget(self.horizontalLayoutWidget_6)
        main_layout.addWidget(self.horizontalLayoutWidget_5)

        # Устанавливаем политику размера для QWidget'ов, содержащих layout'ы.
        # Это важно для корректного масштабирования.
        self.horizontalLayoutWidget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.horizontalLayoutWidget_2.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.horizontalLayoutWidget_3.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.horizontalLayoutWidget_4.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.horizontalLayoutWidget_5.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.horizontalLayoutWidget_6.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)

        # Настройка растягивания элементов внутри горизонтальных layout'ов.
        self.horizontalLayout_2.addStretch(9)
        self.horizontalLayout.setStretch(0, 1)
        self.horizontalLayout.setStretch(1, 1)
        self.horizontalLayout.setStretch(2, 1)
        self.horizontalLayout_2.setStretch(0, 1)
        self.horizontalLayout_2.setStretch(1, 2)  # amount получает больше места
        self.horizontalLayout_2.setStretch(2, 1)

        # Устанавливаем политику размера для amount, чтобы он растягивался
        self.amount.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)

        reg_ex = QRegularExpression("[0-9]*")  # ИЗМЕНЕНО: QRegularExpression
        validator = QRegularExpressionValidator(reg_ex, self.amount)  # ИЗМЕНЕНО: QRegularExpressionValidator
        self.amount.setValidator(validator)

        self.author_label.setStyleSheet("color: rgb(255, 0, 0);")
        self.scs.setStyleSheet("color: green")

        self.amount.setEnabled(False)
        self.amount.setVisible(False)

        self.baton.setEnabled(False)
        self.baton.setVisible(False)

        self.kolvo.setVisible(False)
        self.scs.setVisible(False)


        self.Btn_4.clicked.connect(self.task_4)
        self.Btn_mix.clicked.connect(self.task_19plus)
        self.Btn_10.clicked.connect(self.task_10)
        self.baton.clicked.connect(self.generate)
        self.tool_btn.clicked.connect(self.tool_window)
        self.setWindowTitle('LearnByClick')
        self.words = stress_dict
        self.tool_btn.setText('Редактор заданий')

    def tool_window(self):
        self.tool_wind = tools_w()
        self.tool_wind.exec()

    def generate(self):
        if self.current_task == 4:
            self.create_task_4()
        elif self.current_task == 10:
            self.create_task_10()
        elif self.current_task == 19:
            self.create_task_19plus()

    def task_10(self):
        self.kolvo.setVisible(True)
        self.amount.setEnabled(True)
        self.amount.setVisible(True)
        self.baton.setEnabled(True)
        self.baton.setVisible(True)

        # self.Btn_4.setEnabled(False)
        # self.Btn_4.setVisible(False)
        # self.Btn_mix.setEnabled(False)
        # self.Btn_mix.setVisible(False)

        self.label.setText('Генерация для задания: 10')
        # self.baton.clicked.connect(self.create_task_10)
        self.current_task = 10

    def create_task_10(self):
        k = int(self.amount.text())
        tsk.answers = []
        dc = Document()
        path = self.save_file()
        dc.save(path)
        for i in range(k):
            tsk.fill_10(path, i + 1)
        tsk.fill_doc(tsk.answers, path)
        self.scs.setVisible(True)
        if os.name == 'nt':  # Windows
            os.startfile(path)
        elif os.name == 'posix':  # macOS, Linux
            subprocess.run(['open', path])

    def task_4(self):
        self.kolvo.setVisible(True)
        self.amount.setEnabled(True)
        self.amount.setVisible(True)
        self.baton.setEnabled(True)
        self.baton.setVisible(True)

        # self.Btn_10.setEnabled(False)
        # self.Btn_10.setVisible(False)
        # self.Btn_mix.setEnabled(False)
        # self.Btn_mix.setVisible(False)

        self.label.setText('Генерация для задания: 4')
        self.current_task = 4
        # self.baton.clicked.connect(self.create_task_4)

    def create_task_4(self):
        k = int(self.amount.text())
        tsk.answers = []
        dc = Document()
        path = self.save_file()
        dc.save(path)
        cnt = 1
        for i in range(k):
            correct = tsk.fill_correct()
            incorrect = tsk.fill_incorrect()
            mixed = tsk.fill_mixed()
            answers = tsk.fill_answers(mixed)
            dc = tsk.create_word_doc(mixed, path, cnt)
            dc.save(path)
            cnt += 1
        tsk.fill_doc(tsk.answers, path)
        self.scs.setVisible(True)
        if os.name == 'nt':  # Windows
            os.startfile(path)
        elif os.name == 'posix':  # macOS, Linux
            subprocess.run(['open', path])
        # print(correct, incorrect, mixed, answers, sep='\n')

    def task_19plus(self):
        self.kolvo.setVisible(True)
        self.amount.setEnabled(True)
        self.amount.setVisible(True)
        self.baton.setEnabled(True)
        self.baton.setVisible(True)

        # self.Btn_4.setEnabled(False)
        # self.Btn_4.setVisible(False)
        # self.Btn_10.setEnabled(False)
        # self.Btn_10.setVisible(False)

        self.label.setText('Генерация для задания: 19-21')
        # self.baton.clicked.connect(self.create_task_19plus)
        self.current_task = 19

    def create_task_19plus(self):
        k = int(self.amount.text())
        dc = Document()
        tsk.answers = []
        par = dc.add_paragraph()
        par.add_run('Здесь и далее примеры взяты из Национального корпуса русского языка (ruscorpora.ru)')
        dc.add_paragraph()
        path = self.save_file()
        dc.save(path)
        c = 1
        for i in range(k):
            tsk.fill_19plus(path, c)
            c += 1

        tsk.fill_doc(tsk.answers, path)
        self.scs.setVisible(True)
        if os.name == 'nt':  # Windows
            os.startfile(path)
        elif os.name == 'posix':  # macOS, Linux
            subprocess.run(['open', path])

    def save_file(self):
        file_path, _ = QFileDialog.getSaveFileName(
            caption="Сохранение файла",
            filter="MS Word (*.docx)"
        )
        return file_path

class tools_w(QDialog):
    def __init__(self):
        super().__init__()
        self.initUI()

    def initUI(self):
        self.lay_out = QVBoxLayout()


        self.lay_out.setContentsMargins(100, 120, 100, 120)  # Отступы

        self.setWindowTitle("Редактор заданий")
        self.setGeometry(500, 300, 400, 400)

        # Создаем кнопки
        self.change4 = QPushButton('Редактировать 4 задание', self)
        self.lay_out.addWidget(self.change4)

        self.change10 = QPushButton('Редактировать 10 задание', self)
        self.lay_out.addWidget(self.change10)
        self.change10.clicked.connect(self.change10_behaviour)

        self.change20 = QPushButton('Редактировать 19-21 задание', self)
        self.lay_out.addWidget(self.change20)

        self.guide = QLabel(self)
        self.guide.move(60, 220)
        self.guide.setText('Дополните/удалите строчки таблицы аналогично')
        self.guide.setFixedSize(300, 100)  # Установите фиксированный размер
        self.guide.setVisible(False)

        # растяжение между кнопками
        self.lay_out.addStretch()
        self.setLayout(self.lay_out)

    def change10_behaviour(self):
        editor = Task10Editor(self.guide)
        if self.sender().text() == 'Редактировать 10 задание':
            self.change10.setText('Завершить редактирование')

            self.guide.setVisible(True)
            self.change20.setEnabled(False)
            self.change4.setEnabled(False)

            self.change20.setVisible(False)
            self.change4.setVisible(False)

            editor.get_data()
        else:
            if editor.check_wb():
                if editor.check_data():
                    editor.update_data()

            self.change10.setText('Редактировать 10 задание')

            self.change20.setEnabled(True)
            self.change4.setEnabled(True)

            self.change20.setVisible(True)
            self.change4.setVisible(True)


if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = Main()
    ex.show()
    sys.exit(app.exec())
