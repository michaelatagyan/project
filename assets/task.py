import random
from docx import Document
from .possible_words import stress_dict, comma_list
from .DB_management import DB

class task:
    def __init__(self):
        self.corr = []
        self.incorr = []
        self.words = stress_dict
        self.keys = list(self.words.keys())
        self.answers = []
        self.sentences = comma_list
        self.used_sentences = []

    def fill_correct(self): # Создание списка случайных слов с правильным ударением (4)
        self.corr = []

        i = 0
        self.corr_amount = random.choice(range(1, 5))
        # Заполняем список с правильными словами
        while i != self.corr_amount:
            some_random_word = random.choice(self.keys)

            # Исключаем возможные повторения
            if some_random_word not in self.corr:
                self.corr.append(some_random_word)
                i += 1

        return self.corr

    def fill_incorrect(self): # Создание списка случайных слов с неправильным ударением (4)
        self.incorr = []

        i = 0

        while i != 5 - self.corr_amount:
            some_random_word = self.words[random.choice(self.keys)][0]

            # Исключаем возможные повторения
            if some_random_word not in self.incorr:
                self.incorr.append(some_random_word)
                i += 1

        return self.incorr

    def fill_mixed(self): # Создание списка для записи в файл слов в случайном порядке (4)
        mixed_list = self.corr[:]
        mixed_list.extend(self.incorr)

        # print(mixed_list)

        temp = mixed_list[:]
        mixed_list.clear()
        while temp:
            some_random_word = random.choice(temp)
            mixed_list.append(some_random_word)
            temp.remove(some_random_word)

        return mixed_list

    def fill_answers(self, mix): # Создание списка ответов (4)
        tmp = []
        for b in self.corr:
            if b in mix:
                tmp.append(str(mix.index(b) + 1))

        tmp.sort()
        self.answers.append(tmp)
        return self.answers

    def create_word_doc(self, mix, pt, num): # Запись задания в word-файл (4)
        doc = Document(pt)

        i = 1
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'№{num} '
                          f'Укажите варианты ответов, в которых верно выделена буква, обозначающая ударный гласный'
                          f' звук. Запишите номера ответов.')

        paragraph = doc.add_paragraph()
        for value in mix:
            paragraph.add_run(f'{i}) {value}    ')
            i += 1

        paragraph = doc.add_paragraph()
        paragraph.add_run('Ответ:___________________________')
        paragraph = doc.add_paragraph()
        paragraph.add_run('')
        return doc

    def fill_doc(self, ans, pt): # Запись ответов в файл для проверки (для всех типов заданий одна функция)
        doc = Document(pt)
        doc.add_section()
        i = 1
        for value in ans:
            a = '; '.join(map(str, value))
            paragraph = doc.add_paragraph()
            paragraph.add_run(f'{i}) Ответ: {a}')
            i += 1
        doc.save(pt)

    def fill_19plus(self, pt, m): # Запись задания в word-файл + преобразование строки в формат ... (1) ... (2) ...
        # (19-21)
        doc = Document(pt)
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'№{m}. '
                          f'Укажите цифру(-ы), на месте которой(-ых) должна(-ы) стоять запятая(-ые).')
        paragraph = doc.add_paragraph()

        # Проверка на повторение строки

        sentence = random.choice(self.sentences)
        while True:
            if sentence not in self.used_sentences:
                self.used_sentences.append(sentence)
                break
            else:
                sentence = random.choice(self.sentences)

        # Преобразование строки

        formatted_sentence = ''
        sentence = sentence.split()

        for i in sentence:
            if ',' in i:
                formatted_sentence += f'{i[0:len(i) - 1]} () '
            elif (i.lower() in ['ни', 'а', 'что', 'если', 'хоть', 'то', 'но', 'чтобы', 'пока', 'когда', 'коли', 'или']
                  or 'котор' in i) and formatted_sentence:
                formatted_sentence += f'() {i} '
            elif i.lower() == 'и' or i.lower() == 'однако':
                if formatted_sentence:
                    formatted_sentence += f'() {i} () '
                else:
                    formatted_sentence += f'{i} () '
            else:
                formatted_sentence += i
                formatted_sentence += ' '
        formatted_sentence = formatted_sentence.replace('() ()', '()')
        temp = ''
        n = 1
        for i in formatted_sentence.split():
            if i == '()':
                temp += f'({n}) '
                n += 1
            else:
                temp += i
                temp += ' '
        formatted_sentence = temp
        formatted_sentence = formatted_sentence.strip()

        # Строка преобразована, результат - formatted_sentence

        # print(' '.join(sentence), formatted_sentence, sep='\n')

        paragraph.add_run(formatted_sentence)
        paragraph = doc.add_paragraph()
        paragraph.add_run('Ответ:___________________________')
        paragraph = doc.add_paragraph()
        doc.save(pt)
        self.fill_answers_19plus(sentence, formatted_sentence.split())

    def fill_answers_19plus(self, correct_sentence, formatted_sentence): # Создание списка ответов (19-21)
        n = 0
        answers = []
        for i in range(0, len(formatted_sentence)):
            if '(' in formatted_sentence[i] and ')' in formatted_sentence[i]:
                if correct_sentence[n - 1][-1] == ',':
                    answers.append(formatted_sentence[i][1])
            else:
                n += 1
        self.answers.append(answers)


    def fill_10(self, pt, number):
        doc = Document(pt)
        paragraph = doc.add_paragraph()
        paragraph.add_run(f'№{number}.Укажите варианты ответов, в которых во всех словах одного ряда пропущена одна и та же'
                          f' буква. Запишите номера ответов.')
        paragraph = doc.add_paragraph()

        # Взаимодействие с БД, дабы получить список всех существующих слов
        # database_path = os.path.join(Path(), 'assets', 'identifier.db3')
        # con = sqlite3.connect(database_path)
        # cur = con.cursor()
        data = DB.exec("SELECT * FROM Data")
        words = []
        for i in data:
            words.append(i[1])

        used_rules_ = [] # updated 03.02.2025

        ans = ''
        self.used_sentences.clear()
        while len(ans) not in [1, 2, 3, 4]:
            using = []
            tmp = []
            for _ in range(4): # updated 03.02.2025 (4 вместо 5)
                word = random.choice(words)
                # r = cur.execute("""SELECT * FROM Data WHERE word = ?""", (word,))
                # check_rule = r.fetchall()
                check_rule = DB.exec("SELECT * FROM Data WHERE word=?", word)
                check_rule = check_rule[0][2]
                while check_rule in used_rules_: # updated 03.02.2025
                    word = random.choice(words)
                    # r = cur.execute("""SELECT * FROM Data WHERE word = ?""", (word,)) # updated 03.02.2025
                    # check_rule = r.fetchall()  # updated 03.02.2025
                    check_rule = DB.exec("SELECT * FROM Data WHERE word=?", word)
                    check_rule = check_rule[0][2] # updated 03.02.2025
                used_rules_.append(check_rule) # updated 03.02.2025
                rule = check_rule
                while len(tmp) < 3:
                    if word not in self.used_sentences and rule == check_rule: # added
                        # r = cur.execute("""SELECT * FROM Data WHERE word = ?""", (word,))
                        ap = DB.exec("SELECT * FROM Data WHERE word=?", word)
                        ap = ap[0][3]
                        self.used_sentences.append(word)
                        tmp.append(ap)
                    else:
                        word = random.choice(words)
                        # r = cur.execute("""SELECT * FROM Data WHERE word = ?""", (word,))
                        # rule = r.fetchall()
                        rule = DB.exec("SELECT * FROM Data WHERE word=?", word)
                        rule = rule[0][2]
                using.append(tmp[:])
                tmp.clear()

            used_rules_.clear() # updated 03.02.2025
            c = 0
            for j in using:
                seq = ''
                for i in j:
                    # r = cur.execute("""SELECT * FROM Data WHERE formatted = ?""", (i,))
                    # ltr = r.fetchall()
                    ltr = DB.exec("SELECT * FROM Data WHERE formatted=?", i)
                    ltr = ltr[0][-1]
                    seq += ltr
                c += 1
                if seq[0] == seq[1] == seq[2]:
                    ans += str(c)
        self.answers.append(ans)

        c = 1
        for i in using:
            line = f'{c}) '
            for j in i:
                line += j
                line += ', '
            c += 1
            line = line[0:len(line) - 2]
            paragraph.add_run(line)
            paragraph = doc.add_paragraph()
        paragraph.add_run('Ответ:___________________________')
        paragraph = doc.add_paragraph()
        doc.save(pt)
        # con.close()


tsk = task()
